# For reading file name
import glob
# попробуем работать с подушкой
from PIL import Image
from datetime import datetime
from docx import Document
from docx.shared import Cm
import os

print(glob.glob('.'))  # распечатаем все названия файлов в директории

myListWithoutSlash = []  # создадим список из фотографий
for file_name in glob.iglob('./imgs/*.*', recursive=True):
    print(file_name)
    file_name = file_name.split('\\')[1]  # убираем слэши в названии и берем только последнее
    myListWithoutSlash.append(file_name)
document = Document()  # create word doc
table = document.add_table(rows=1, cols=3)  # добавляе таблицу
table.style = 'Table Grid'
# give names to columns
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Номер'
hdr_cells[1].text = 'Описание'
hdr_cells[2].text = 'Фотография'
n = 1  # счетчик
for i in range(len(myListWithoutSlash)):
    print(i)
    row_cells = table.add_row().cells  # добавляем строку к табилце
    # lets make row height 11 cm
    row_cells[0].text = f"фото {n}"  # содержание первого столбца
    picture = Image.open(f'./imgs/{myListWithoutSlash[i]}')  # open image
    width, height = picture.size # get picture size in pixels
    text = ''
    flag = False
    t = 0
    for s in range(len(myListWithoutSlash[i])):
        g = myListWithoutSlash[i][s - 1 - 2 * t]
        if myListWithoutSlash[i][s - 1 - 2 * t] == '.':
            flag = True
        if flag:
            text += myListWithoutSlash[i][s - 1 - 2 * t]
        t += 1
    text = text[::-1] # reverse text to get correct name
    picture.save(f'./imgs/{text}_{n}.jpeg')  # пересохраняем в jpeg
    row_cells[1].text = text  # содержание второго столбца
    p = row_cells[2].add_paragraph()
    r = p.add_run()
    # lets make all pictures in doc have the same width
    # so height will be calculated with regard to height / width ratio
    picture_constant = height / width
    r.add_picture(f'./imgs/{text}_{n}.jpeg', width=Cm(
                 13),
                 height=Cm(13 * picture_constant))
    picture.close()
    os.remove(f'./imgs/{myListWithoutSlash[i]}')  # удаляем старое фото
    n += 1
# go through all rows and change its height to 11 cm instead of first one
for row in range(len(table.rows)):
    if row == 0:
        pass
    else:
        table.rows[row].height = Cm(11)
# change doc margins
sections = document.sections
margin = 0.5
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)
# для названия файле используем datetime
id_cur = datetime.now().strftime('%d_%m_%y__%H%M')
document.save(f'./output/готов_{id_cur}.docx')
